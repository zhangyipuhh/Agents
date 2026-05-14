from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import os
import re
from app.shared.utils.report.word.config import ReportConfig, SectionConfig, CoverElementConfig, HeadingStyleConfig, ParagraphStyleConfig
logger = logging.getLogger(__name__)

# 兜底：确保所有日志同时输出到控制台，绕过logging handler配置问题
_orig_info = logger.info
_orig_debug = logger.debug

def _info(msg, *args, **kwargs):
    print(f"[WordReport] {msg}")
    _orig_info(msg, *args, **kwargs)

def _debug(msg, *args, **kwargs):
    print(f"[WordReport.DEBUG] {msg}")
    _orig_debug(msg, *args, **kwargs)

logger.info = _info
logger.debug = _debug




def set_chinese_font(run, font_name="宋体", font_size=12, bold=False):
    """
    设置中文字体的辅助函数

    Args:
        run: python-docx的Run对象，用于设置文本格式
        font_name: 字体名称，默认为"宋体"
        font_size: 字体大小（磅值），默认为12
        bold: 是否加粗，默认为False

    Returns:
        无返回值，直接修改run对象的字体属性
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold


def set_first_line_indent_chars(paragraph, char_count: int):
    """
    设置首行缩进（字符单位）

    通过直接操作XML设置w:ind元素的w:firstLineChars属性，
    实现以字符为单位的首行缩进，缩进量随字体大小自动调整。

    Args:
        paragraph: python-docx的Paragraph对象，将被修改首行缩进属性
        char_count: 缩进字符数，如2表示缩进2个中文字符

    Returns:
        无返回值，直接修改段落的XML属性

    Notes:
        - w:firstLineChars的值 = 字符数 × 100，即200=2字符
        - Word会根据当前字体自动计算对应的w:firstLine值（twips单位）
        - 相比厘米单位(w:firstLine)，字符单位能随字体大小自动调整缩进量
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(char_count * 100))


def set_line_spacing(paragraph, line_value: float, line_rule: str = "auto"):
    """
    设置段落行间距

    通过直接操作XML设置w:spacing元素的w:line和w:lineRule属性，
    支持固定行距、最小行距和多倍行距三种模式。

    Args:
        paragraph: python-docx的Paragraph对象，将被修改行间距属性
        line_value: 行距值，根据line_rule不同含义不同：
            - "auto"模式：float类型，表示行距倍数（如1.5=1.5倍行距）
            - "exact"模式：float类型，表示固定行距的磅值（如28=28磅）
            - "atLeast"模式：float类型，表示最小行距的磅值（如20=20磅）
        line_rule: 行距规则，可选值：
            - "auto": 多倍行距（w:lineRule="auto"）
            - "exact": 固定行距（w:lineRule="exact"）
            - "atLeast": 最小行距（w:lineRule="atLeast"）

    Returns:
        无返回值，直接修改段落的XML属性

    Raises:
        ValueError: 当line_rule不是"auto"/"exact"/"atLeast"之一时

    Notes:
        - "auto"模式下，w:line的值 = 倍数 × 240（240 twips = 12pt = 单倍行距基准）
        - "exact"和"atLeast"模式下，w:line的值 = 磅值 × 20（1pt = 20 twips）
        - w:line的单位始终是缇(twips)
    """
    valid_rules = ("auto", "exact", "atLeast")
    if line_rule not in valid_rules:
        raise ValueError(f"line_rule 必须是 {valid_rules} 之一，得到: {line_rule}")

    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)

    if line_rule == "auto":
        twips_value = int(line_value * 240)
    else:
        twips_value = int(line_value * 20)

    spacing.set(qn('w:line'), str(twips_value))
    spacing.set(qn('w:lineRule'), line_rule)


class WordReportGenerator:
    """
    Word报告通用生成器

    通过ReportConfig配置对象驱动生成Word文档，支持封面、目录、
    多级标题、正文段落、分页等元素的灵活组合，支持自动生成目录页码和页脚页码。
    """

    def __init__(self, config: ReportConfig):
        """
        初始化报告生成器

        Args:
            config: ReportConfig配置对象，定义报告的结构和内容
        """
        self.config = config
        self.doc = None
        self._heading_bookmarks = []
        self._heading_index = 0
        self._bookmark_id_counter = 0

    def _add_field_code(self, paragraph, field_code: str, font_name: str = "宋体", font_size: int = 10):
        """
        在段落中添加Word域字段代码

        Args:
            paragraph: python-docx的Paragraph对象，域字段将作为段落的直接子元素添加
            field_code: 域字段代码，如"PAGE"、"NUMPAGES"、"PAGEREF bookmark_name"
            font_name: 字体名称，默认"宋体"
            font_size: 字号，单位pt，默认10

        Notes:
            使用w:fldChar元素创建Word域字段，实现页码等自动计算功能。
            标准域字段结构：每个部分（begin/instrText/separate/display/end）
            作为独立的w:r元素直接添加到段落w:p中，避免w:r嵌套导致Word解析异常。

            正确的XML结构：
            <w:p>
              <w:r><w:rPr>...</w:rPr><w:fldChar type="begin"/></w:r>
              <w:r><w:rPr>...</w:rPr><w:instrText> PAGE </w:instrText></w:r>
              <w:r><w:rPr>...</w:rPr><w:fldChar type="separate"/></w:r>
              <w:r><w:rPr>...</w:rPr><w:t>1</w:t></w:r>
              <w:r><w:rPr>...</w:rPr><w:fldChar type="end"/></w:r>
             </w:p>
        """
        logger.debug(f"添加域字段: {field_code}")
        def _make_run():
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(font_size * 2))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(font_size * 2))
            rPr.append(szCs)
            run.append(rPr)
            return run

        run_begin = _make_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_begin.set(qn('w:dirty'), 'true')
        run_begin.append(fldChar_begin)
        paragraph._p.append(run_begin)

        run_instr = _make_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' {field_code} \\* MERGEFORMAT \\h '
        run_instr.append(instrText)
        paragraph._p.append(run_instr)

        run_separate = _make_run()
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        run_separate.append(fldChar_separate)
        paragraph._p.append(run_separate)

        run_display = _make_run()
        display_t = OxmlElement('w:t')
        display_t.text = '1'
        run_display.append(display_t)
        paragraph._p.append(run_display)

        run_end = _make_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run_end.append(fldChar_end)
        paragraph._p.append(run_end)

    def _add_bookmark(self, paragraph, bookmark_name: str):
        """
        为段落添加书签

        Args:
            paragraph: python-docx的Paragraph对象
            bookmark_name: 书签名称

        Returns:
            str: 书签名称

        Notes:
            使用w:bookmarkStart和w:bookmarkEnd元素创建Word书签
        """
        # 生成唯一书签ID
        bookmark_id = str(self._bookmark_id_counter)
        self._bookmark_id_counter += 1
        logger.debug(f"添加书签: {bookmark_name} id={bookmark_id}")

        # 创建bookmarkStart元素
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id)
        bookmark_start.set(qn('w:name'), bookmark_name)

        # 创建bookmarkEnd元素
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_id)

        # 添加到段落的首个run之前和最后一个run之后
        if paragraph._p.getchildren():
            first_run = paragraph._p.find(qn('w:r'))
            if first_run is not None:
                paragraph._p.insert(paragraph._p.index(first_run), bookmark_start)
            else:
                paragraph._p.append(bookmark_start)
        else:
            paragraph._p.append(bookmark_start)

        paragraph._p.append(bookmark_end)

        return bookmark_name

    def _pre_collect_bookmarks(self):
        """
        预收集标题书签名称

        在渲染目录之前，扫描config.sections中所有heading类型的section，
        按顺序生成书签名称并填充_heading_bookmarks列表，
        使目录渲染时能正确引用标题书签生成PAGEREF域字段。

        Notes:
            - 书签命名规则：_TocHeading_{index}，与_render_heading中的命名一致
            - 必须在_render_toc()之前调用
            - 预收集确保目录中的PAGEREF域字段能正确引用尚未渲染的标题书签
        """
        for section in self.config.sections:
            if section.section_type == "heading":
                bookmark_name = f"_TocHeading_{len(self._heading_bookmarks)}"
                self._heading_bookmarks.append(bookmark_name)
        logger.info(f"预收集完成，共{len(self._heading_bookmarks)}个标题书签")

    def _setup_page(self):
        """
        设置页面尺寸和边距

        根据ReportConfig中page_setup的配置，设置文档所有节的页面宽度、高度和四边边距。

        Notes:
            - 所有尺寸单位为厘米(cm)，通过Cm()转换为docx内部单位
            - 遍历文档所有section，确保每个节的页面参数一致
        """
        ps = self.config.page_setup
        for section in self.doc.sections:
            section.page_width = Cm(ps.page_width)
            section.page_height = Cm(ps.page_height)
            section.top_margin = Cm(ps.top_margin)
            section.bottom_margin = Cm(ps.bottom_margin)
            section.left_margin = Cm(ps.left_margin)
            section.right_margin = Cm(ps.right_margin)

    def _setup_default_style(self):
        """
        设置默认字体样式

        将文档Normal样式的字体设置为config中指定的默认字体名称和大小，
        同时设置东亚字体(w:eastAsia)以确保中文字符正确显示。

        Notes:
            - 修改Normal样式会影响文档中所有未单独指定字体的段落
            - rPr.rFonts.set(qn('w:eastAsia'), ...) 用于设置东亚文字字体
        """
        style = self.doc.styles['Normal']
        style.font.name = self.config.default_font_name
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.default_font_name)
        style.font.size = Pt(self.config.default_font_size)

    def _resolve_text(self, content):
        """
        解析文本内容，支持变量替换和可调用对象

        Args:
            content: 文本内容，可以是字符串（支持{{变量名}}占位符）或可调用对象

        Returns:
            str: 解析后的文本内容

        Notes:
            - 字符串中的{{变量名}}会被config.data中对应的值替换
            - 可调用对象会被调用，传入config.data作为参数，返回值作为文本
            - 如果content为None，返回空字符串
        """
        if callable(content):
            return str(content(self.config.data))
        if isinstance(content, str) and self.config.data:
            def replace_var(match):
                var_name = match.group(1)
                return str(self.config.data.get(var_name, match.group(0)))
            return re.sub(r'\{\{(.+?)\}\}', replace_var, content)
        return str(content) if content is not None else ""

    def _render_cover_element(self, element: CoverElementConfig):
        """
        渲染单个封面元素

        Args:
            element: CoverElementConfig配置对象，包含文本、字体、对齐和间距信息

        Notes:
            - 根据alignment设置对齐方式（left/center/right）
            - 应用space_before和space_after控制段前段后行数
            - 使用指定的字体名称、大小和加粗样式
        """
        if not element or not element.text:
            return

        # 段前行数
        for _ in range(element.space_before):
            self.doc.add_paragraph()

        # 创建段落
        p = self.doc.add_paragraph()

        # 设置对齐方式
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = alignment_map.get(element.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        # 添加文本
        run = p.add_run(self._resolve_text(element.text))
        set_chinese_font(run, element.font_name, element.font_size, element.bold)

        # 段后行数
        for _ in range(element.space_after):
            self.doc.add_paragraph()

    def _render_cover(self):
        """
        渲染封面页

        根据CoverConfig配置生成封面，包含附件编号、标题、副标题、日期等元素。
        封面结束后自动添加新节（分节符），使封面成为独立的Section。

        Notes:
            - 附件编号：默认右对齐显示
            - 标题：默认居中显示，使用封面专用字体和字号
            - 副标题：默认居中显示
            - 日期：默认居中显示
            - 各元素通过space_before和space_after控制段前段后行数
            - 如果cover配置为None，则跳过封面渲染
            - 封面结束后使用add_section创建新节，便于独立控制页脚页码
        """
        cover = self.config.cover
        if not cover:
            return

        # 按顺序渲染各元素：附件编号 -> 主标题 -> 副标题 -> 日期
        self._render_cover_element(cover.attachment)
        self._render_cover_element(cover.title)
        self._render_cover_element(cover.subtitle)
        self._render_cover_element(cover.date)

        # 封面结束，添加新节（新页）
        self.doc.add_section(WD_SECTION_START.NEW_PAGE)

    def _render_toc(self):
        """
        渲染目录页

        根据TocConfig配置生成目录，支持多级缩进和自动页码。
        目录结束后自动添加新节（分节符），使目录成为独立的Section。

        Notes:
            - 目录标题：居中显示，使用目录专用字体和字号
            - 目录条目：根据level级别设置左缩进，实现层级效果
            - 条目字体：优先使用TocEntry.font_name，为空时回退到TocConfig.entry_font_name
            - 条目加粗：优先使用TocEntry.bold，为None时回退到TocConfig.entry_bold
            - 缩进量 = indent_per_level * level（单位cm）
            - 页码：通过Word域字段自动生成，使用制表位右对齐
            - 引导符：标题与页码之间使用leader指定的字符连接
            - 如果toc配置为None，则跳过目录渲染
            - 目录结束后使用add_section创建新节，便于独立控制页脚页码
        """
        toc = self.config.toc
        if not toc:
            return

        # 目录标题（居中）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(toc.title)
        set_chinese_font(run, toc.title_font_name, toc.title_font_size, toc.title_bold)

        # 空行
        self.doc.add_paragraph()

        # 目录条目
        logger.info(f"开始渲染目录，共{len(toc.entries)}个条目")
        for idx, entry in enumerate(toc.entries):
            p = self.doc.add_paragraph()

            # 设置制表位
            tab_stops = p.paragraph_format.tab_stops
            # 添加右对齐制表位，带引导符
            leader_map = {
                "...": 1,  # 点号引导
                "---": 2,  # 虚线引导
                "___": 3,  # 下划线引导
            }
            leader = leader_map.get(toc.leader, 0)  # 默认无引导符

            # 计算页码位置（页面宽度 - 右边距 - 左边距）
            page_width = self.config.page_setup.page_width
            left_margin = self.config.page_setup.left_margin
            right_margin = self.config.page_setup.right_margin
            tab_position = page_width - left_margin - right_margin - 1  # 留出1cm边距

            from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
            tab_stops.add_tab_stop(Cm(tab_position), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS if leader == 1 else WD_TAB_LEADER.NONE)

            # 添加标题文本
            entry_font_name = entry.font_name if entry.font_name else toc.entry_font_name
            entry_font_size = entry.font_size if entry.font_size is not None else toc.entry_font_size
            entry_bold = entry.bold if entry.bold is not None else toc.entry_bold
            run = p.add_run(self._resolve_text(entry.text))
            set_chinese_font(run, entry_font_name, entry_font_size, entry_bold)

            # 添加制表符
            p.add_run("\t")

            # 添加页码（使用域字段）
            if toc.show_page_number and self._heading_bookmarks:
                # 条目数可能多于标题数，超出部分使用最后一个书签
                if idx < len(self._heading_bookmarks):
                    bookmark_name = self._heading_bookmarks[idx]
                else:
                    bookmark_name = self._heading_bookmarks[-1]
                self._add_field_code(
                    p, f"PAGEREF {bookmark_name}",
                    toc.page_number_font_name, toc.page_number_font_size
                )

            # 设置左缩进
            p.paragraph_format.left_indent = Cm(toc.indent_per_level * entry.level)

        # 目录结束，添加新节（新页）
        self.doc.add_section(WD_SECTION_START.NEW_PAGE)
        logger.info("目录渲染完成")

    def _render_heading(self, section: SectionConfig):
        """
        渲染标题段落

        Args:
            section: SectionConfig配置对象，包含标题的级别、文本和样式信息

        Notes:
            样式优先级：SectionConfig显式设置 > ReportConfig.heading_styles[level]统一样式 > 内置默认值
            - 从heading_styles获取对应级别的统一样式配置
            - SectionConfig中的font_name/font_size/bold/space_before/space_after/left_indent
              非空/非零时可覆盖统一样式配置
            - 自动添加书签，用于目录页码引用
        """
        level = section.level
        text = self._resolve_text(section.content)

        heading_style: HeadingStyleConfig = self.config.heading_styles.get(level, HeadingStyleConfig())

        font_name = section.font_name or heading_style.font_name
        font_size = section.font_size or heading_style.font_size
        bold = section.bold if section.bold is not None else heading_style.bold

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        set_chinese_font(run, font_name, font_size, bold)

        space_before = section.space_before if section.space_before is not None else heading_style.space_before
        paragraph.paragraph_format.space_before = Pt(space_before)

        space_after = section.space_after if section.space_after is not None else heading_style.space_after
        paragraph.paragraph_format.space_after = Pt(space_after)

        left_indent = section.left_indent if section.left_indent is not None else heading_style.left_indent
        if left_indent:
            paragraph.paragraph_format.left_indent = Cm(left_indent)

        if section.alignment is not None:
            paragraph.alignment = section.alignment

        bookmark_name = self._heading_bookmarks[self._heading_index]
        logger.debug(f"渲染标题 Lv.{level}: {text} → 书签: {bookmark_name}")
        self._add_bookmark(paragraph, bookmark_name)
        self._heading_index += 1

    def _render_paragraph(self, section: SectionConfig):
        """
        渲染正文段落

        Args:
            section: SectionConfig配置对象，包含段落的文本和样式信息

        Notes:
            样式优先级：SectionConfig显式设置 > ReportConfig.paragraph_style统一样式 > 内置默认值
            - 从paragraph_style获取正文段落的统一样式配置
            - SectionConfig中的font_name/font_size/bold/first_line_indent/line_spacing_rule/
              line_spacing_value/space_before/space_after/left_indent
              非空/非零时可覆盖统一样式配置
            - 首行缩进使用字符单位(w:firstLineChars)，缩进量随字体大小自动调整
            - 行间距支持三种模式：auto(多倍行距)/exact(固定行距)/atLeast(最小行距)
        """
        text = self._resolve_text(section.content)

        para_style: ParagraphStyleConfig = self.config.paragraph_style

        font_name = section.font_name or para_style.font_name
        font_size = section.font_size or para_style.font_size
        bold = section.bold if section.bold is not None else para_style.bold

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        set_chinese_font(run, font_name, font_size, bold)

        if section.alignment is not None:
            paragraph.alignment = section.alignment
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        first_indent = section.first_line_indent if section.first_line_indent is not None else para_style.first_line_indent_chars
        if first_indent:
            set_first_line_indent_chars(paragraph, first_indent)

        line_rule = section.line_spacing_rule or para_style.line_spacing_rule
        line_value = section.line_spacing_value if section.line_spacing_value is not None else para_style.line_spacing_value
        set_line_spacing(paragraph, line_value, line_rule)

        space_before = section.space_before if section.space_before is not None else para_style.space_before
        if space_before:
            paragraph.paragraph_format.space_before = Pt(space_before)

        space_after = section.space_after if section.space_after is not None else para_style.space_after
        if space_after:
            paragraph.paragraph_format.space_after = Pt(space_after)

        left_indent = section.left_indent if section.left_indent is not None else para_style.left_indent
        if left_indent:
            paragraph.paragraph_format.left_indent = Cm(left_indent)

    def _render_page_break(self):
        """
        渲染分页符

        在当前位置插入一个分页符，后续内容将从新页面开始。

        Notes:
            - 分页符不会产生可见的段落内容
            - 常用于封面后、目录后等需要强制换页的场景
        """
        self.doc.add_page_break()

    def _set_update_fields(self):
        """
        设置文档打开时自动更新域字段

        通过在settings.xml中添加w:updateFields元素，
        使Word打开文档时自动更新所有域字段（包括目录中的PAGEREF页码），
        确保目录页码在打开文档后立即正确显示。

        Notes:
            - 添加w:updateFields val="true"到文档设置中
            - 用户首次打开文档时Word会提示是否更新域，确认后目录页码自动刷新
            - 此设置对PAGEREF、PAGE、NUMPAGES等所有域字段生效
        """
        settings = self.doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            update_fields.set(qn('w:val'), 'true')
            # 按OOXML Schema顺序，updateFields须在decimalSymbol之前
            # 直接append到末尾会导致WPS/部分Word版本忽略该元素
            decimal_symbol = settings.find(qn('w:decimalSymbol'))
            if decimal_symbol is not None:
                settings.insert(list(settings).index(decimal_symbol), update_fields)
                logger.debug("updateFields 已插入到 decimalSymbol 之前")
            else:
                settings.append(update_fields)
                logger.debug("updateFields 已追加到 settings 末尾")

    def _set_start_page(self, section, start_page: int):
        """
        设置Section的起始页码

        Args:
            section: python-docx的Section对象
            start_page: 起始页码数值

        Notes:
            通过设置w:pgNumType元素的w:start属性来指定节的起始页码。
            如果w:pgNumType元素不存在则自动创建。
        """
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), str(start_page))

    def _render_footer(self):
        """
        渲染页脚页码

        根据FooterConfig配置，在文档各节的页脚中添加自动页码。
        根据Section角色（封面/目录/正文）和start_from配置分别处理页码显示逻辑。

        Notes:
            - 使用Word域字段PAGE和NUMPAGES实现自动页码
            - 支持{page}和{total}占位符
            - 封面Section：根据skip_cover决定是否显示页码
            - 目录Section：根据skip_toc决定是否显示页码
            - 正文Section：始终显示页码
            - start_from配置决定页码从哪个节开始计数（cover/toc/content）
            - 在start_from指定的节上设置起始页码，后续节延续计数
            - 通过_add_field_code方法添加Word域字段
        """
        footer_config = self.config.footer
        if not footer_config or not footer_config.enabled:
            return

        sections = self.doc.sections

        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        alignment = alignment_map.get(footer_config.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        has_cover = self.config.cover is not None
        has_toc = self.config.toc is not None

        section_roles = []
        if has_cover:
            section_roles.append("cover")
        if has_toc:
            section_roles.append("toc")
        while len(section_roles) < len(sections):
            section_roles.append("content")

        start_from = footer_config.start_from
        start_page_set = False

        for idx, section in enumerate(sections):
            role = section_roles[idx] if idx < len(section_roles) else "content"
            footer = section.footer
            footer.is_linked_to_previous = False

            for paragraph in footer.paragraphs:
                paragraph.clear()

            if not start_page_set and role == start_from:
                self._set_start_page(section, footer_config.start_page)
                start_page_set = True

            skip_page_number = False
            if role == "cover" and footer_config.skip_cover:
                skip_page_number = True
            elif role == "toc" and footer_config.skip_toc:
                skip_page_number = True

            if skip_page_number:
                continue

            if not footer.paragraphs:
                footer.add_paragraph()

            paragraph = footer.paragraphs[0]
            paragraph.alignment = alignment

            format_str = footer_config.format
            parts = re.split(r'(\{page\}|\{total\})', format_str)

            for part in parts:
                if part == "{page}":
                    self._add_field_code(
                        paragraph, "PAGE",
                        footer_config.font_name, footer_config.font_size
                    )
                elif part == "{total}":
                    self._add_field_code(
                        paragraph, "NUMPAGES",
                        footer_config.font_name, footer_config.font_size
                    )
                else:
                    run = paragraph.add_run(part)
                    set_chinese_font(run, footer_config.font_name, footer_config.font_size)
        logger.info(f"页脚页码已渲染，共{len(sections)}个节，起始节: {start_from}")

    def _render_section(self, section: SectionConfig):
        """
        根据段落类型分发渲染

        Args:
            section: SectionConfig配置对象

        Notes:
            根据 section_type 字段分发到对应的渲染方法：
            - "heading" -> _render_heading()
            - "paragraph" -> _render_paragraph()
            - "page_break" -> _render_page_break()
        """
        if section.section_type == "heading":
            self._render_heading(section)
        elif section.section_type == "paragraph":
            self._render_paragraph(section)
        elif section.section_type == "page_break":
            self._render_page_break()

    def generate(self) -> Document:
        """
        生成Word报告的主入口方法

        按顺序执行：创建文档 -> 默认样式 -> 封面 -> 目录 -> 正文段落 -> 页面设置 -> 页脚页码 -> 更新域设置

        Returns:
            Document: python-docx的Document对象，调用方可通过 doc.save() 保存为.docx文件

        Notes:
            - 每次调用generate()都会创建新的Document对象
            - 如果cover或toc配置为None，对应部分会被跳过
            - sections列表中的每个SectionConfig按顺序依次渲染
            - 标题自动添加书签，用于目录页码引用
            - 页脚页码通过Word域字段自动生成
            - 页面设置在所有节创建完成后执行，确保每个节的页面参数一致
            - 设置w:updateFields使Word打开文档时自动更新目录页码等域字段
        """
        logger.info("开始生成Word报告")
        self.doc = Document()
        self._heading_bookmarks = []
        self._heading_index = 0
        self._bookmark_id_counter = 0
        self._setup_default_style()
        self._render_cover()
        logger.info("封面渲染完成")
        self._pre_collect_bookmarks()
        self._render_toc()

        for section in self.config.sections:
            self._render_section(section)
        logger.info(f"正文内容渲染完成，共{len(self.config.sections)}个章节")

        self._setup_page()
        logger.info("页面设置完成")
        self._render_footer()
        self._set_update_fields()
        logger.info("Word报告生成完成")

        return self.doc

    def save(self, file_path: str):
        """
        保存Word文档到指定路径

        Args:
            file_path: 输出文件路径，如 "output/report.docx"

        Raises:
            ValueError: 如果尚未调用 generate() 生成文档

        Notes:
            - 调用前必须先调用generate()方法生成文档
            - 文件路径的目录必须存在，否则会抛出异常
            - 文件格式为.docx（Office Open XML格式）
        """
        if self.doc is None:
            raise ValueError("请先调用 generate() 方法生成文档")
        self.doc.save(file_path)
        logger.info(f"文档已保存到: {file_path}")
        self._update_fields_via_com(file_path)

    def _update_fields_via_com(self, file_path: str):
        """
        使用Word COM自动化更新文档中所有域字段

        保存后的docx中PAGEREF等域字段显示为默认值，通过Word COM打开文档、
        更新所有域字段后保存，使目录页码显示为实际值。

        Args:
            file_path: 已保存的.docx文件绝对路径

        Notes:
            - 仅在Windows上可用，需要安装Microsoft Word或WPS Office
            - 若未安装win32com或COM初始化失败，则跳过更新并记录warning日志
            - Word COM会短暂启动Word进程（不可见），更新完成后自动关闭
        """
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            abs_path = os.path.abspath(file_path)
            logger.info(f"Word COM: 正在更新域字段...")
            doc = word.Documents.Open(abs_path)
            doc.Fields.Update()
            doc.Save()
            doc.Close()
            word.Quit()
            logger.info("Word COM: 域字段更新完成")
        except ImportError:
            logger.warning("未安装 win32com 模块，跳过域字段自动更新")
        except Exception as e:
            logger.warning(f"Word COM 更新域字段失败: {e}")
