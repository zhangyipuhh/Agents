#!/usr/bin/python
# -*- coding:utf-8 -*-
"""
WordLoader 模块

提供 Word 文件加载功能，基于 python-docx 进行封装，
支持延迟加载、文件存在性检查等扩展功能。
支持读取段落文本、表格文本，以及合同文档的特殊处理。

Date: 2026-03-13
Author: 张镒谱
"""

from pathlib import Path
from typing import Optional, List, Dict, Tuple, Any
from docx import Document as DocxDocument
from langchain_core.documents import Document


class WordLoader:
    """
    Word 文件加载器
    
    对 python-docx 进行封装，提供更友好的 API 和额外功能。
    使用延迟加载模式优化性能，只有在实际调用 load 方法时才创建底层文档对象。
    
    Attributes:
        file_path: 要加载的 Word 文件路径
        load_method: 加载方法，可选 "default", "replace", "contract", "paragraphs"
        pattern: 正则表达式模式（用于 replace 和 contract 方法）
        pattern_replace: 替换后的文本（用于 replace 和 contract 方法）
        paragraph_num: 要加载的段落数量（用于 paragraphs 方法）
    """
    
    def __init__(
        self,
        file_path: str,
        load_method: str = "default",
        pattern: Optional[str] = None,
        pattern_replace: Optional[str] = None,
        paragraph_num: Optional[int] = None
    ):
        """
        初始化 WordLoader
        
        Args:
            file_path: 要加载的 Word 文件路径，支持相对路径和绝对路径
            load_method: 加载方法，可选 "default", "replace", "contract", "paragraphs"
            pattern: 正则表达式模式（用于 replace 和 contract 方法）
            pattern_replace: 替换后的文本（用于 replace 和 contract 方法）
            paragraph_num: 要加载的段落数量（用于 paragraphs 方法）
        """
        self.file_path = Path(file_path)
        self.load_method = load_method
        self.pattern = pattern
        self.pattern_replace = pattern_replace
        self.paragraph_num = paragraph_num
        self._doc: Optional[DocxDocument] = None

    def _get_document(self) -> DocxDocument:
        """
        获取底层的 python-docx Document 实例
        
        使用延迟加载模式，只有在首次调用时才创建文档实例，
        避免不必要的资源占用。
        
        Returns:
            DocxDocument: 底层的 python-docx 文档实例
        """
        if self._doc is None:
            self._doc = DocxDocument(str(self.file_path))
        return self._doc

    def load(self) -> List[Document]:
        """
        加载 Word 文件内容
        
        根据 load_method 参数调用不同的加载方法：
        - "default": 默认加载所有段落和表格
        - "replace": 使用正则表达式替换文本后加载
        - "contract": 加载合同文档，返回段落数据
        - "paragraphs": 只加载前 N 个段落
        
        Returns:
            List[Document]: Document 对象列表
        """
        if self.load_method == "replace":
            if self.pattern is None or self.pattern_replace is None:
                raise ValueError("使用 'replace' 方法时需要提供 pattern 和 pattern_replace 参数")
            return self.load_with_replace(self.pattern, self.pattern_replace)
        elif self.load_method == "contract":
            if self.pattern is None or self.pattern_replace is None:
                raise ValueError("使用 'contract' 方法时需要提供 pattern 和 pattern_replace 参数")
            contract_text, _ = self.load_contract(self.pattern, self.pattern_replace)
            # 将合同文本包装为 Document 列表
            return [Document(
                page_content=contract_text,
                metadata={
                    "source": str(self.file_path),
                    "type": "contract",
                    "method": "contract"
                }
            )]
        elif self.load_method == "paragraphs":
            if self.paragraph_num is None:
                raise ValueError("使用 'paragraphs' 方法时需要提供 paragraph_num 参数")
            return self.load_by_paragraphs(self.paragraph_num)
        else:
            # 默认加载方法
            return self._load_default()

    def _load_default(self) -> List[Document]:
        """
        默认加载 Word 文件内容
        
        将 Word 文件内容转换为 Document 对象列表，
        每个 Document 包含一段文本内容及其元数据。
        
        Returns:
            List[Document]: Document 对象列表
        """
        doc = self._get_document()
        documents = []
        
        # 提取段落文本
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                documents.append(
                    Document(
                        page_content=para.text,
                        metadata={
                            "source": str(self.file_path),
                            "type": "paragraph",
                            "index": idx
                        }
                    )
                )
        
        # 提取表格文本
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                cell_content = " | ".join(row_text)
                if cell_content.strip():
                    documents.append(
                        Document(
                            page_content=cell_content,
                            metadata={
                                "source": str(self.file_path),
                                "type": "table",
                                "table_index": table_idx,
                                "row_index": row_idx
                            }
                        )
                    )
        
        return documents

    def lazy_load(self) -> List[Document]:
        """
        延迟加载 Word 文件内容
        
        与 load 方法类似，但使用 lazy_load 模式，适合处理大文件。
        当前实现与 load 相同，后续可优化为真正的生成器模式。
        
        Returns:
            List[Document]: Document 对象列表
        """
        return self.load()

    def load_with_replace(
        self, 
        pattern: str, 
        pattern_replace: str
    ) -> List[Document]:
        """
        加载 Word 文件内容，并替换指定文本
        
        Args:
            pattern: 要替换的正则表达式模式
            pattern_replace: 替换后的文本
            
        Returns:
            List[Document]: Document 对象列表
        """
        import re
        
        doc = self._get_document()
        documents = []
        
        # 提取段落文本并替换
        for idx, para in enumerate(doc.paragraphs):
            text = re.sub(pattern, pattern_replace, para.text)
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": str(self.file_path),
                            "type": "paragraph",
                            "index": idx,
                            "replaced": True
                        }
                    )
                )
        
        # 提取表格文本并替换
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    text = re.sub(pattern, pattern_replace, cell.text)
                    row_text.append(text)
                cell_content = " | ".join(row_text)
                if cell_content.strip():
                    documents.append(
                        Document(
                            page_content=cell_content,
                            metadata={
                                "source": str(self.file_path),
                                "type": "table",
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "replaced": True
                            }
                        )
                    )
        
        return documents

    def load_contract(
        self,
        pattern: str,
        pattern_replace: str
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        专门加载制式合同 Word 文档内容

        合同格式固定且没有表格，如果有表格这个暂时搞不定。
        会返回整个合同文本和段落数据结构。

        Args:
            pattern: 要替换的正则表达式模式
            pattern_replace: 替换后的文本

        Returns:
            Tuple[str, List[Dict]]: 整个合同文本和段落数据列表
            段落数据格式示例:
            [
                {
                    "paragraph_type": 1,  # 第几条款,比如1、2、3、4
                    "paragraph_num": 1,   # 段落的实际索引
                    "paragraph_text": "段落1文本"
                },
                {
                    "paragraph_type": 1,
                    "paragraph_num": 2,
                    "paragraph_text": "段落2文本"
                }
            ]
        """
        import re

        doc = self._get_document()
        paragraph_data = []
        contract_lines = []

        _index = 0
        _tmp_index = 0

        for para in doc.paragraphs:
            text = re.sub(pattern, pattern_replace, para.text)

            # 检查是否匹配条款模式
            _is_match = re.match(pattern, text)
            if _is_match:
                _tmp_index += 1

            # 只记录非空段落
            if text.strip():
                paragraph_data.append({
                    "paragraph_type": _tmp_index,
                    "paragraph_num": _index,
                    "paragraph_text": text
                })
                contract_lines.append(text)
                _index += 1
            else:
                _index += 1

        contract_text = "\n".join(contract_lines)
        return contract_text, paragraph_data

    def load_by_paragraphs(self, paragraph_num: int) -> List[Document]:
        """
        加载 Word 文档前 N 个段落
        
        Args:
            paragraph_num: 要加载的段落数量
            
        Returns:
            List[Document]: Document 对象列表
        """
        doc = self._get_document()
        documents = []
        index = 0
        
        for para in doc.paragraphs:
            if para.text.strip() == "":
                continue
            if index >= paragraph_num:
                break
                
            documents.append(
                Document(
                    page_content=para.text,
                    metadata={
                        "source": str(self.file_path),
                        "type": "paragraph",
                        "index": index
                    }
                )
            )
            index += 1
        
        return documents

    @property
    def exists(self) -> bool:
        """
        检查文件是否存在
        
        Returns:
            bool: 文件存在返回 True，否则返回 False
        """
        return self.file_path.exists()


if __name__ == "__main__":
    # 测试代码
    # 方式1: 使用 load_contract 方法直接获取合同文本和段落数据
    loader = WordLoader("D:\DocumentLoader\WordLoader.docx")
    if loader.exists:
        contract_text, paragraph_data = loader.load_contract(r'^\s*(第[\u4e00-\u9fa5\d]+条)', r'\1条款')
        print("合同文本:")
        print(contract_text)
        print("\n段落数据:")
        print(paragraph_data)

    # 方式2: 使用参数配置方式加载合同
    loader2 = WordLoader(
        "D:\DocumentLoader\WordLoader.docx",
        load_method="contract",
        pattern=r'^\s*(第[\u4e00-\u9fa5\d]+条)',
        pattern_replace=r'\1条款'
    )
    if loader2.exists:
        documents = loader2.load()  # 这会返回 Document 列表
        print("\n通过 load() 加载合同:")
        for doc in documents:
            print(doc.page_content)
