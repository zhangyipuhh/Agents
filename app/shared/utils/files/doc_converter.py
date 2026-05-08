#!/usr/bin/python
# -*- coding:utf-8 -*-
"""
文档格式转换工具 - 将旧版 .doc 格式转换为 .docx 格式

用于支持知识库中的 .doc 文件预览功能。

Date: 2026/05/08
"""

import os
import io
import tempfile
import logging
import subprocess
import shutil
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


def find_libreoffice_path() -> Optional[str]:
    """
    在系统中查找 LibreOffice 可执行文件路径
    
    Returns:
        str: soffice.exe 的完整路径，如果未找到则返回 None
    """
    possible_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files\LibreOffice 8\program\soffice.exe",
        r"C:\Users\{user}\AppData\Local\LibreOffice\Program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    ]
    
    for path in possible_paths:
        if "{user}" in path:
            try:
                user_home = os.path.expanduser("~")
                path = path.replace("{user}", os.path.basename(user_home))
            except:
                continue
        
        if os.path.exists(path):
            logger.info(f"找到 LibreOffice: {path}")
            return path
    
    try:
        result = subprocess.run(
            ["where", "soffice"] if os.name == 'nt' else ["which", "soffice"],
            capture_output=True,
            text=True,
            timeout=10
        )
        if result.returncode == 0:
            found_path = result.stdout.strip().split('\n')[0]
            if os.path.exists(found_path):
                logger.info(f"从 PATH 找到 LibreOffice: {found_path}")
                return found_path
    except:
        pass
    
    try:
        import winreg
        key = winreg.OpenKey(
            winreg.HKEY_LOCAL_MACHINE,
            r"SOFTWARE\LibreOffice\Components"
        )
        install_path, _ = winreg.QueryValueEx(key, "InstallPath")
        soffice_path = os.path.join(install_path, "program", "soffice.exe")
        if os.path.exists(soffice_path):
            logger.info(f"从注册表找到 LibreOffice: {soffice_path}")
            return soffice_path
    except:
        pass
    
    return None

def convert_doc_to_docx(doc_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    将 .doc 文件转换为 .docx 格式

    Args:
        doc_path: .doc 文件的完整路径

    Returns:
        Tuple[转换后的文件路径, 错误信息]
        成功时: (docx_path, None)
        失败时: (None, error_message)
    """
    if not os.path.exists(doc_path):
        return None, f"文件不存在: {doc_path}"

    if not doc_path.lower().endswith('.doc'):
        return None, "文件不是 .doc 格式"

    temp_dir = tempfile.gettempdir()
    base_name = os.path.splitext(os.path.basename(doc_path))[0]
    docx_path = os.path.join(temp_dir, f"{base_name}_converted_{os.getpid()}.docx")

    try:
        try:
            import win32com.client
            result = _convert_using_word(doc_path, docx_path)
            if result[0]:
                return result
        except (ImportError, Exception) as e:
            logger.warning(f"Word 转换失败: {e}，尝试其他方式")

        result = _convert_using_libreoffice(doc_path, docx_path)
        if result[0]:
            return result

        result = _convert_using_mammoth(doc_path, docx_path)
        if result[0]:
            return result

        return None, "所有转换方式均失败：未安装 Word、LibreOffice 或 mammoth"

    except Exception as e:
        logger.error(f"转换失败: {e}")
        return None, f"转换失败: {str(e)}"


def _convert_using_word(doc_path: str, output_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    使用 Microsoft Word COM 对象转换文档（仅 Windows）

    Args:
        doc_path: 源文件路径
        output_path: 输出文件路径

    Returns:
        Tuple[转换后的文件路径, 错误信息]
    """
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs(os.path.abspath(output_path), FileFormat=16)
        doc.Close()
        word.Quit()

        if os.path.exists(output_path):
            return output_path, None
        else:
            return None, "Word 转换后文件未生成"

    except ImportError as e:
        logger.warning(f"pywin32 未安装: {e}")
        return None, "pywin32 未安装"
    except Exception as e:
        logger.error(f"Word COM 转换失败: {e}")
        return None, f"Word 转换失败: {str(e)}"


def _convert_using_libreoffice(doc_path: str, output_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    使用 LibreOffice 转换文档

    Args:
        doc_path: 源文件路径
        output_path: 输出文件路径

    Returns:
        Tuple[转换后的文件路径, 错误信息]
    """
    soffice_path = find_libreoffice_path()

    if not soffice_path:
        return None, "未找到 LibreOffice，无法转换"

    temp_dir = tempfile.mkdtemp()
    base_name = os.path.splitext(os.path.basename(doc_path))[0]
    temp_doc_path = os.path.join(temp_dir, f"{base_name}.doc")

    try:
        shutil.copy(doc_path, temp_doc_path)

        filter_options = [
            "MS Word 97",
            "doc8",
            "Text Encoded RTF w/oansinew"
        ]
        
        last_error = None
        for filter_name in filter_options:
            cmd = [
                soffice_path,
                "--headless",
                "--infilter=" + filter_name,
                "--convert-to", "docx",
                "--outdir", temp_dir,
                temp_doc_path
            ]
            
            logger.info(f"尝试 LibreOffice 转换 (filter: {filter_name}): {' '.join(cmd)}")
            
            try:
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=180
                )
                
                if result.returncode == 0:
                    converted_file = os.path.join(temp_dir, f"{base_name}.docx")
                    if os.path.exists(converted_file):
                        shutil.copy(converted_file, output_path)
                        logger.info(f"LibreOffice 转换成功 (filter: {filter_name})")
                        return output_path, None
                
                last_error = result.stderr if result.stderr else f"Filter {filter_name} failed"
                logger.warning(f"Filter {filter_name} 失败: {last_error}")
                
            except subprocess.TimeoutExpired:
                last_error = "LibreOffice 转换超时 (180秒)"
                logger.warning(last_error)
                continue
            except Exception as e:
                last_error = f"LibreOffice 执行异常: {str(e)}"
                logger.error(last_error)
                continue

        return None, f"LibreOffice 转换失败: {last_error}"

    except Exception as e:
        logger.error(f"LibreOffice 转换异常: {e}")
        return None, f"LibreOffice 转换异常: {str(e)}"
    finally:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except:
            pass


def get_libreoffice_installation_guide() -> str:
    """
    返回 LibreOffice 安装指引
    
    Returns:
        str: 安装指引文本
    """
    return """
    请安装 LibreOffice 以支持 .doc 文件转换:
    
    下载地址: https://www.libreoffice.org/download/download/
    
    Windows 安装步骤:
    1. 访问上述地址下载最新版本
    2. 运行安装程序，按默认选项安装
    3. 安装完成后重启应用程序
    4. 验证: 运行 check_conversion_support() 检查 libreoffice 是否为 True
    
    Linux 安装命令:
    sudo apt install libreoffice-writer
    
    Mac 安装:
    brew install --cask libreoffice
    """


def _convert_using_mammoth(doc_path: str, output_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    使用 mammoth 库将 .doc 文件转换为 .docx 格式（纯 Python 实现）

    Args:
        doc_path: 源文件路径
        output_path: 输出文件路径

    Returns:
        Tuple[转换后的文件路径, 错误信息]
    """
    try:
        import mammoth
        from docx import Document
        import re

        with open(doc_path, 'rb') as doc_file:
            result = mammoth.convert_to_html(doc_file)
            html_content = result.value

        doc = Document()

        clean_html = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        clean_html = re.sub(r'<script[^>]*>.*?</script>', '', clean_html, flags=re.DOTALL | re.IGNORECASE)

        blocks = re.split(r'<(?:p|div|li|tr|h[1-6])[^>]*>', clean_html, flags=re.IGNORECASE)

        for block in blocks:
            if not block.strip():
                continue

            clean_text = re.sub(r'<[^>]+>', '', block)
            clean_text = clean_text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>')
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()

            if clean_text:
                doc.add_paragraph(clean_text)

        doc.save(output_path)

        if os.path.exists(output_path):
            logger.info(f"mammoth 转换成功: {doc_path} -> {output_path}")
            return output_path, None
        else:
            return None, "mammoth 转换后文件未生成"

    except ImportError:
        return None, "mammoth 未安装，请运行: pip install mammoth python-docx"
    except Exception as e:
        logger.error(f"mammoth 转换异常: {e}")
        return None, f"mammoth 转换异常: {str(e)}"


def create_simple_docx_from_text(text_content: str, output_path: str) -> Tuple[bool, Optional[str]]:
    """
    从纯文本内容创建简单的 .docx 文件
    这是一个兜底方案，当无法转换时使用

    Args:
        text_content: 文本内容
        output_path: 输出文件路径

    Returns:
        Tuple[是否成功, 错误信息]
    """
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        for line in text_content.split('\n'):
            para = doc.add_paragraph()
            para.add_run(line)

        doc.save(output_path)
        return True, None

    except ImportError:
        return False, "python-docx 未安装"
    except Exception as e:
        logger.error(f"创建 docx 失败: {e}")
        return False, str(e)


def check_conversion_support() -> dict:
    """
    检查系统支持的转换方式

    Returns:
        dict: 支持情况
    """
    support = {
        "pywin32": False,
        "libreoffice": False,
        "libreoffice_path": None,
        "python_docx": False
    }

    try:
        import win32com.client
        support["pywin32"] = True
    except ImportError:
        pass

    libreoffice_path = find_libreoffice_path()
    if libreoffice_path:
        support["libreoffice"] = True
        support["libreoffice_path"] = libreoffice_path
        logger.info(f"LibreOffice 支持: {libreoffice_path}")
    else:
        logger.warning("LibreOffice 未安装或未找到")

    try:
        from docx import Document
        support["python_docx"] = True
    except ImportError:
        pass

    return support


if __name__ == "__main__":
    support = check_conversion_support()
    print("转换支持情况:")
    for key, value in support.items():
        print(f"  {key}: {'支持' if value else '不支持'}")