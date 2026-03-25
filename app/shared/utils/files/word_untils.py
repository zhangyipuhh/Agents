"""
Word处理助手 - 负责处理Word文档，包括读取和修改
"""
import os
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
from collections import defaultdict
import io
import re
import itertools
import copy

import docx
from docx import Document
from docx.text.run import Run
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.core.config.config import WORD_OUTPUT_CONFIG

def isolate_run(paragraph, start, end):
    """
    将段落中指定范围的文本隔离到一个独立的run中
    这是python-docx社区的标准解决方案，用于处理跨run的文本标记

    Args:
        paragraph: docx段落对象
        start: 文本起始位置（在段落文本中的索引）
        end: 文本结束位置（在段落文本中的索引）

    Returns:
        Run: 包含目标文本的独立run对象
    """
    runs = list(paragraph.runs)

    if not runs:
        raise ValueError("段落中没有runs")

    # 计算每个run的累积结束位置
    r_ends = list(itertools.accumulate(len(r.text) for r in runs))

    # 定位包含start的run
    r_idx = 0
    for i, r_end in enumerate(r_ends):
        if start < r_end:
            r_idx = i
            break
    else:
        r_idx = len(r_ends) - 1

    current_run = runs[r_idx]
    skipped = r_ends[r_idx - 1] if r_idx > 0 else 0
    start_in_r = start - skipped
    end_in_r = end - skipped

    # 拆分前缀（如果start不在run开头）
    if start_in_r > 0:
        prefix_text = current_run.text[:start_in_r]
        prefix_run = copy.deepcopy(current_run)
        prefix_run.text = prefix_text
        current_run.text = current_run.text[start_in_r:]
        current_run._r.addprevious(prefix_run._r)
        end_in_r -= start_in_r

    # 处理后缀或合并后续runs
    if len(current_run.text) > end_in_r:
        suffix_text = current_run.text[end_in_r:]
        current_run.text = current_run.text[:end_in_r]
        suffix_run = copy.deepcopy(current_run)
        suffix_run.text = suffix_text
        current_run._r.addnext(suffix_run._r)
    elif len(current_run.text) < end_in_r:
        while len(current_run.text) < end_in_r:
            needed = end_in_r - len(current_run.text)
            r_idx += 1
            if r_idx >= len(runs):
                break
            next_run = runs[r_idx]
            if len(next_run.text) <= needed:
                current_run.text += next_run.text
                next_run._r.getparent().remove(next_run._r)
            else:
                current_run.text += next_run.text[:needed]
                next_run.text = next_run.text[needed:]

    return current_run


class WordProcessor:
    """Word处理器类"""
    
    def __init__(self):
        """初始化Word处理器"""
        pass
    def read_contract_word(self, file_path: str,pattern:str,pattern_replace: str) -> Tuple[str,dict]:
        """
        专门读取制式合同Word文档内容，合同格式固定且没有表格，如果有表格这个暂时搞不定
        
        Args:
            file_path: Word文档路径
            pattern: 要替换的文本
            pattern_replace: 替换后的文本
        Returns:
            文档文本内容,段落数据
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # 提取段落文本，在提取段落文本时，如果is_replace_words为True，则替换文本
            #每个段落标记一个序号，生成一个字典
            '''
            格式示例
            [
                {
                    "paragraph_type": "标题",#第几条款,比如1、2、3、4
                    "paragraph_num": 1,#段落的实际索引
                    "paragraph_text": "段落1文本"
                },
                {
                    "paragraph_type": "正文",#第几条款
                    "paragraph_num": 2,#段落的实际索引
                    "paragraph_text": "段落2文本"
                }
            ]
            '''
            _index = 0
            _paragraph_data = []
            _tmp_match = ""#比如都是第四条，type都写一样的，直到遇到新的条款
            #为空的行不记录，但是index会加1
            _tmp_index = 0#条款索引
            for para in doc.paragraphs:
                #para.text = para.text.replace(" ", "")
                para.text = re.sub(pattern,pattern_replace,para.text)
                full_text.append(para.text)
                _is_match = re.match(pattern, para.text)
                if _is_match:
                    _tmp_index += 1
                if para.text.strip():
                    
                    _paragraph_data.append({
                        "paragraph_type":_tmp_index,
                        "paragraph_num": _index,
                        "paragraph_text": para.text
                    })
                    _index += 1
                else:
                    _index += 1
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text),_paragraph_data
        except Exception as e:
            print(f"读取Word文档时出错: {str(e)}")
            return ""
    def read_word(self, file_path: str,is_replace_words: bool = False,pattern: str = None,pattern_replace: str = None) -> str:
        """
        读取Word文档内容
        
        Args:
            file_path: Word文档路径
            is_replace_words: 是否替换文本
            pattern: 要替换的文本
            pattern_replace: 替换后的文本
        Returns:
            文档文本内容
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if is_replace_words and pattern and pattern_replace:
                    para.text = re.sub(pattern,pattern_replace,para.text)
                full_text.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取Word文档时出错: {str(e)}")
            return ""
    def read_word_by_paragraphs(self, file_path: str,paragraph_num: int) -> str:
        """
        读取Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档文本内容
        """
        try:
            doc = Document(file_path)
            full_text = []
            index = 0
            # 提取段落文本
            for para in doc.paragraphs :
                if para.text.strip() == "":
                    continue
                if index > paragraph_num:
                    break
                full_text.append(para.text)
                index += 1
            
            
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取Word文档时出错: {str(e)}")
            return ""

    def highlight_text(self, doc_path: str, text_to_highlight: str, output_path: str = None) -> str:
        """
        在Word文档中高亮指定文本
        
        Args:
            doc_path: Word文档路径
            text_to_highlight: 要高亮的文本
            output_path: 输出文档路径，如果为None则使用默认路径
            
        Returns:
            输出文档的路径
        """
        if output_path is None:
            # 创建输出目录
            os.makedirs(WORD_OUTPUT_CONFIG["output_dir"], exist_ok=True)
            
            # 生成输出文件名
            file_name = os.path.basename(doc_path)
            base_name, ext = os.path.splitext(file_name)
            output_path = os.path.join(WORD_OUTPUT_CONFIG["output_dir"], f"{base_name}_标记{ext}")
        
        try:
            # 加载文档
            doc = Document(doc_path)
            
            # 高亮颜色
            highlight_color = WORD_OUTPUT_CONFIG["highlight_color"]
            color = RGBColor.from_string(highlight_color)
            
            # 编译正则表达式用于匹配文本
            pattern = re.compile(re.escape(text_to_highlight), re.IGNORECASE)
            
            # 遍历段落并高亮匹配文本
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text_to_highlight in text:
                    matches = list(pattern.finditer(text))
                    for match in matches:
                        start, end = match.span()
                        try:
                            target_run = isolate_run(paragraph, start, end)
                            target_run.font.color.rgb = color
                        except Exception as e:
                            print(f"高亮文本时出错: {e}")
                            continue

            # 遍历表格单元格并高亮匹配文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            if text_to_highlight in text:
                                matches = list(pattern.finditer(text))
                                for match in matches:
                                    start, end = match.span()
                                    try:
                                        target_run = isolate_run(paragraph, start, end)
                                        target_run.font.color.rgb = color
                                    except Exception as e:
                                        print(f"高亮文本时出错: {e}")
                                        continue
            
            # 保存文档
            doc.save(output_path)
            print(f"已将标记结果保存至: {output_path}")
            return output_path
        
        except Exception as e:
            print(f"高亮文本时出错: {str(e)}")
            return doc_path
    def highlight_text_without_changing_format(self, doc_path, text_to_find, color=(255, 0, 0), output_path=None, only_first_occurrence=False):
        """
        在Word文档中查找并高亮文本，保留原始格式
        
        参数:
            doc_path: Word文档路径
            text_to_find: 要查找的文本，可以是逗号分隔的多个文本
            color: RGB颜色元组，默认为红色 (255, 0, 0)，用于背景高亮
            output_path: 输出文档路径，如果为None则覆盖原文件
            only_first_occurrence: 是否只标记第一次出现的位置，默认为False
        """
        # 加载文档
        doc = Document(doc_path)
    
        color_hex = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
        
        # 分割text_to_find为列表，处理可能的引号和空格问题
        texts_to_find = []
        for item in text_to_find.split(','):
            item = item.strip()
            # 移除可能的引号
            if (item.startswith('"') and item.endswith('"')) or (item.startswith("'") and item.endswith("'")):
                item = item[1:-1]
            if item:  # 确保不添加空字符串
                texts_to_find.append(item)
        
        # 删除可能的重复项
        texts_to_find = list(dict.fromkeys(texts_to_find))
        
        # 打印调试信息
        print(f"待标记的文本项: {texts_to_find}")
        
        # 全局计数器，记录每个待查找文本已标记的次数
        text_marked_count = {text: 0 for text in texts_to_find}
        
        # 定义应用高亮样式的函数
        def apply_highlighting_to_run(run, color_hex):
            """为一个run应用高亮样式"""
            try:
                rPr = run._r.get_or_add_rPr()
                
                # 删除现有的阴影元素
                for shd in rPr.findall(qn("w:shd")):
                    rPr.remove(shd)
                
                # 创建新的阴影元素
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), color_hex)
                rPr.append(shd)
                return True
            except Exception as e:
                print(f"应用高亮样式时出错: {str(e)}")
                return False
            
        # 检查一个run是否完全匹配查找文本
        def is_exact_match(run_text, text_to_highlight):
            return run_text == text_to_highlight
            
        # 查找段落中的文本并高亮
        def find_and_highlight_in_paragraph(paragraph, text_to_highlight):
            """在段落中查找文本并高亮，不改变段落结构"""
            if not text_to_highlight or not paragraph.text:
                return False
            
            # 快速检查：如果段落不包含要查找的文本，直接返回
            if text_to_highlight not in paragraph.text:
                return False
            
            # 检查是否已全局标记过且只需标记一次
            if only_first_occurrence and text_marked_count[text_to_highlight] > 0:
                return False
                
            highlighted = False
            
            # 检查每个run是否精确匹配查找文本
            for run in paragraph.runs:
                # 如果run文本精确匹配查找文本
                if is_exact_match(run.text, text_to_highlight):
                    if apply_highlighting_to_run(run, color_hex):
                        highlighted = True
                        text_marked_count[text_to_highlight] += 1
                        if only_first_occurrence:
                            return True
            
            # 如果没有精确匹配，则查找部分匹配并精确隔离
            if not highlighted:
                # 构建段落完整文本并定位每个run
                full_text = ""
                run_positions = []
                for i, run in enumerate(paragraph.runs):
                    start_pos = len(full_text)
                    full_text += run.text
                    end_pos = len(full_text)
                    run_positions.append((start_pos, end_pos, run))

                # 查找所有匹配位置
                start_idx = 0
                while True:
                    idx = full_text.find(text_to_highlight, start_idx)
                    if idx == -1:
                        break

                    match_end = idx + len(text_to_highlight)

                    # 使用isolate_run精确隔离目标文本，然后应用高亮
                    try:
                        target_run = isolate_run(paragraph, idx, match_end)
                        if apply_highlighting_to_run(target_run, color_hex):
                            highlighted = True
                            text_marked_count[text_to_highlight] += 1
                            if only_first_occurrence:
                                return True
                    except Exception as e:
                        print(f"隔离文本时出错: {e}")

                    start_idx = idx + 1
            
            return highlighted
        
        # 收集所有段落（主文档和表格）
        all_paragraphs = []
        
        # 收集主文档段落
        for paragraph in doc.paragraphs:
            all_paragraphs.append(paragraph)
        
        # 收集表格段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_paragraphs.append(paragraph)
        
        # 处理每个要查找的文本
        for text_to_find in texts_to_find:
            # 重置计数器
            text_marked_count[text_to_find] = 0
            
            # 处理所有段落
            for paragraph in all_paragraphs:
                try:
                    # 如果已经标记过且只需标记一次，则跳过
                    if only_first_occurrence and text_marked_count[text_to_find] > 0:
                        break
                        
                    find_and_highlight_in_paragraph(paragraph, text_to_find)
                except Exception as e:
                    print(f"处理段落时出错: {str(e)}")
                    continue
        
        # 保存文档
        try:
            if output_path:
                doc.save(output_path)
            else:
                doc.save(doc_path)
            
            return output_path if output_path else doc_path
        except Exception as e:
            print(f"保存文档时出错: {str(e)}")
            return doc_path

    def highlight_text_in_specific_paragraph(
        self, 
        doc_path: str, 
        text_to_find: Union[str, List[str]], 
        paragraph_indices: Union[int, List[int]], 
        color: Tuple[int, int, int] = (255, 0, 0), 
        output_path: str = None
    ) -> str:
        """
        在Word文档的指定段落中批量标记特定文本，文本和段落一一对应
        
        Args:
            doc_path: Word文档路径
            text_to_find: 要标记的文本，可以是单个字符串或字符串列表
            paragraph_indices: 要标记的段落索引（从0开始），可以是单个整数或整数列表，
                             必须与text_to_find一一对应
            color: RGB颜色元组，默认为红色 (255, 0, 0)
            output_path: 输出文档路径，如果为None则覆盖原文件
            
        Returns:
            输出文档的路径
            
        Examples:
            # 单个文本和段落
            processor.highlight_text_in_specific_paragraph("doc.docx", "要标记的文本", 1)
            
            # 多个文本在对应的段落（一一对应）
            processor.highlight_text_in_specific_paragraph(
                "doc.docx", 
                ["文本1", "文本2"], # 文本1在第1段，文本2在第2段
                [1, 2]
            )
        """
        try:
            # 加载文档
            doc = Document(doc_path)
            
            # 转换颜色为十六进制
            color_hex = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
            
            # 将输入转换为列表形式
            texts = [text_to_find] if isinstance(text_to_find, str) else text_to_find
            indices = [paragraph_indices] if isinstance(paragraph_indices, int) else paragraph_indices
            
            # 参数验证
            if not texts or not indices:
                print("文本列表或段落索引列表不能为空")
                return doc_path
                
            # 确保文本和段落索引数量相同
            if len(texts) != len(indices):
                print(f"文本数量({len(texts)})与段落索引数量({len(indices)})不匹配")
                return doc_path
                
            # 验证段落索引的有效性，如果超出范围会提醒
            total_paragraphs = len(doc.paragraphs)
            # 检查每个段落索引数组中的索引是否有效
            invalid_indices = []
            for idx_array in indices:
                # 如果是单个索引,转换为列表方便统一处理
                if isinstance(idx_array, int):
                    idx_array = [idx_array]
                    
                # 检查数组中每个索引的有效性
                for idx in idx_array:
                    if idx < 0 or idx >= total_paragraphs:
                        invalid_indices.append(idx)
            #invalid_indices = [idx for idx in indices if idx < 0 or idx >= total_paragraphs]
            if invalid_indices:
                print(f"发现无效的段落索引 {invalid_indices}。文档共有 {total_paragraphs} 个段落")
                return doc_path
            
            # 对每对文本和段落进行处理
            for text,para_idx in zip(texts,indices):
                # 构建段落完整文本并定位每个run
                full_text = ""
                run_positions = []
                
                #将整个条款都拼到一起
                for para_idx_tmp in para_idx:
                    paragraph = doc.paragraphs[para_idx_tmp]
                    for run in paragraph.runs:
                        start_pos = len(full_text)
                        full_text += run.text
                        end_pos = len(full_text)
                        run_positions.append((start_pos, end_pos, run))
                
                
                
                # 如果段落中不包含目标文本，跳过此段落
                if text not in full_text:
                    #如果没有尝试去掉左右空格，因为原文档如果不标准，会存在大量空格
                    if text.replace(" ", "") not in full_text.replace(" ", ""):
                        print(f"未找到文本: {text}")
                        continue
                    
                # 查找所有匹配位置
                start_idx = 0
                while True:
                    idx = full_text.find(text, start_idx)
                    if idx == -1:
                        # 去掉空格后查找
                        stripped_full = full_text.replace(" ", "")
                        stripped_text = text.replace(" ", "")
                        idx = stripped_full.find(stripped_text, start_idx)

                        if idx == -1:
                            break

                        # 将去掉空格后的索引映射回原始文本（考虑空格数量）
                        space_count = 0
                        real_idx = 0
                        for i, char in enumerate(full_text):
                            if char == ' ':
                                space_count += 1
                            if space_count > idx:
                                real_idx = i
                                break
                            real_idx = i + 1
                        idx = real_idx
                    
                    # 找出匹配落在哪个段落的哪个位置
                    match_end = idx + len(text)
                    para_start = 0
                    target_paragraph = None
                    target_start = 0
                    target_end = 0

                    for para_idx_tmp in para_idx:
                        paragraph = doc.paragraphs[para_idx_tmp]
                        para_len = len(paragraph.text)
                        para_end = para_start + para_len

                        # 检查匹配是否在这个段落内
                        if idx >= para_start and match_end <= para_end:
                            target_paragraph = paragraph
                            target_start = idx - para_start
                            target_end = match_end - para_start
                            break
                        para_start = para_end

                    if target_paragraph is None:
                        # 如果没找到精确匹配，尝试找到起始段落
                        for para_idx_tmp in para_idx:
                            paragraph = doc.paragraphs[para_idx_tmp]
                            para_len = len(paragraph.text)
                            para_end = para_start + para_len
                            if idx < para_end:
                                target_paragraph = paragraph
                                target_start = max(0, idx - para_start)
                                target_end = min(para_len, match_end - para_start)
                                break
                            para_start = para_end

                    if target_paragraph:
                        try:
                            # 使用isolate_run精确隔离目标文本
                            target_element = isolate_run(target_paragraph, target_start, target_end)
                            # 应用背景高亮
                            rPr = target_element.get_or_add_rPr()
                            for shd in rPr.findall(qn("w:shd")):
                                rPr.remove(shd)
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:val"), "clear")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:fill"), color_hex)
                            rPr.append(shd)
                        except Exception as e:
                            print(f"隔离文本时出错: {e}")
                    
                    start_idx = idx + 1  # 继续搜索下一个匹配项
            
            # 保存文档
            save_path = output_path if output_path else doc_path
            doc.save(save_path)
            print(f"文档已保存至: {save_path}")
            return save_path
            
        except Exception as e:
            print(f"标记文本时出错: {str(e)}")
            return doc_path

    def mark_contract_dynamic_values(
        self,
        doc_path: str,
        marking_rules: List[Dict[str, Any]],
        output_path: str = None
    ) -> str:
        """
        根据标记规则自动识别并标记合同中的动态值
        
        Args:
            doc_path: Word文档路径
            marking_rules: 标记规则列表，每个规则包含：
                - clause_pattern: 条款编号模式（如"第六条"、"第五条"）
                - prefix_pattern: 动态值前缀
                - suffix_pattern: 动态值后缀（可选）
                - color: RGB颜色元组
            output_path: 输出文档路径，如果为None则覆盖原文件
        
        Returns:
            输出文档的路径
        
        Raises:
            IOError: 文件操作失败时抛出
            ValueError: 参数验证失败时抛出
        
        Examples:
            marking_rules = [
                {
                    "clause_pattern": "第六条",
                    "prefix_pattern": "用途为",
                    "suffix_pattern": None,
                    "color": (255, 0, 0)
                },
                {
                    "clause_pattern": "第五条",
                    "prefix_pattern": "大写",
                    "suffix_pattern": "（小写",
                    "color": (255, 0, 0)
                }
            ]
            processor.mark_contract_dynamic_values("contract.docx", marking_rules)
        """
        # 参数验证：确保文档路径存在
        if not os.path.exists(doc_path):
            raise IOError(f"文档路径不存在: {doc_path}")
        
        # 参数验证：确保标记规则不为空
        if not marking_rules:
            raise ValueError("标记规则不能为空")
        
        try:
            # 加载文档
            doc = Document(doc_path)
            
            # 处理每条标记规则
            for rule in marking_rules:
                clause_pattern = rule.get("clause_pattern")
                prefix_pattern = rule.get("prefix_pattern")
                suffix_pattern = rule.get("suffix_pattern")
                color = rule.get("color", (255, 0, 0))
                
                # 验证规则必要字段
                if not clause_pattern or not prefix_pattern:
                    print(f"警告: 规则缺少必要字段，跳过该规则: {rule}")
                    continue
                
                # 转换颜色为十六进制格式
                color_hex = f"{color[0]:02x}{color[1]:02x}{color[2]:02x}"
                
                # 查找条款所在的段落索引
                clause_paragraph_indices = self._find_clause_paragraphs(doc, clause_pattern)
                
                # 如果未找到条款，记录警告并跳过
                if not clause_paragraph_indices:
                    print(f"警告: 未找到条款 '{clause_pattern}'，跳过该规则")
                    continue
                
                success = self._mark_dynamic_value_in_clause(
                    doc, clause_paragraph_indices, prefix_pattern, suffix_pattern, color_hex
                )
                
                if not success:
                    print(f"[兜底] 在条款 '{clause_pattern}' 中未找到动态值，尝试标红条款标题")
                    self._mark_clause_title(doc, clause_paragraph_indices, clause_pattern, color_hex)
            
            # 确定保存路径
            save_path = output_path if output_path else doc_path
            
            # 保存文档
            doc.save(save_path)
            print(f"文档已保存至: {save_path}")
            return save_path
            
        except Exception as e:
            raise IOError(f"处理文档时出错: {str(e)}")
    
    def _get_punctuation_variants(self, pattern: str) -> List[str]:
        """
        获取标点符号的中英文变体列表
        
        Args:
            pattern: 原始模式字符串
        
        Returns:
            包含中英文变体的列表
        """
        punctuation_map = {
            ',': ['，', ','],
            '，': ['，', ','],
            '。': ['。', '.'],
            '.': ['。', '.'],
            ';': ['；', ';'],
            '；': ['；', ';'],
            ':': ['：', ':'],
            '：': ['：', ':'],
            '(': ['（', '('],
            '（': ['（', '('],
            ')': ['）', ')'],
            '）': ['）', ')'],
            '"': ['"', '"', '"'],
            '"': ['"', '"', '"'],
            '"': ['"', '"', '"'],
        }
        
        variants = [pattern]
        for eng, chinese_variants in punctuation_map.items():
            if eng in pattern or chinese_variants[0] in pattern:
                for v in chinese_variants:
                    if eng in pattern:
                        new_variant = pattern.replace(eng, v)
                    else:
                        new_variant = pattern.replace(chinese_variants[0], v)
                    if new_variant not in variants:
                        variants.append(new_variant)
        
        return variants
    
    def _find_clause_paragraphs(self, doc: Document, clause_pattern: str) -> List[int]:
        """
        查找包含指定条款编号的所有段落索引
        
        Args:
            doc: Document对象
            clause_pattern: 条款编号模式（如"第六条"）
        
        Returns:
            包含该条款的段落索引列表
        
        Note:
            条款可能跨越多个段落，此方法返回从条款开始到下一个条款之间的所有段落索引
        """
        clause_indices = []
        found_clause = False
        
        clause_regex = re.compile(clause_pattern)
        next_clause_regex = re.compile(r'第[\u4e00-\u9fa5\d]+条')
        
        for i, paragraph in enumerate(doc.paragraphs):
            if found_clause:
                if next_clause_regex.search(paragraph.text) and not clause_regex.search(paragraph.text):
                    break
                clause_indices.append(i)
            elif clause_regex.search(paragraph.text):
                found_clause = True
                clause_indices.append(i)
        
        return clause_indices
    
    def _mark_dynamic_value_in_clause(
        self,
        doc: Document,
        clause_paragraph_indices: List[int],
        prefix_pattern: str,
        suffix_pattern: Optional[str],
        color_hex: str
    ) -> bool:
        """
        在指定条款段落中标记动态值

        Args:
            doc: Document对象
            clause_paragraph_indices: 条款段落索引列表
            prefix_pattern: 动态值前缀
            suffix_pattern: 动态值后缀（可选）
            color_hex: 十六进制颜色值

        Returns:
            是否成功标记
        """
        # 标准化文本用于匹配（移除所有空白字符）
        def normalize_text(text):
            return ''.join(text.split())

        for para_idx in clause_paragraph_indices:
            paragraph = doc.paragraphs[para_idx]
            para_text = paragraph.text
            
            print(f"[调试] 检查段落 {para_idx}: {para_text[:80]}...")
            print(f"[调试] 查找前缀: '{prefix_pattern}'")

            if prefix_pattern not in para_text and normalize_text(prefix_pattern) not in normalize_text(para_text):
                print(f"[调试] 段落 {para_idx} 中未找到前缀 '{prefix_pattern}'")
                continue

            print(f"[调试] 在段落 {para_idx} 中找到前缀模式 '{prefix_pattern}'")
            
            full_text = ""
            run_positions = []
            
            for run in paragraph.runs:
                start_pos = len(full_text)
                full_text += run.text
                end_pos = len(full_text)
                run_positions.append((start_pos, end_pos, run))
            
            # 查找前缀位置，支持多种空白字符变体
            prefix_idx = full_text.find(prefix_pattern)
            if prefix_idx == -1:
                # 尝试标准化匹配（移除所有空白字符）
                normalized_full = normalize_text(full_text)
                normalized_prefix = normalize_text(prefix_pattern)
                normalized_idx = normalized_full.find(normalized_prefix)
                if normalized_idx == -1:
                    print(f"[调试] 在完整文本中未找到前缀 '{prefix_pattern}'")
                    continue
                # 将标准化索引映射回原始文本索引
                prefix_idx = 0
                char_count = 0
                for i, char in enumerate(full_text):
                    if not char.isspace():
                        if char_count == normalized_idx:
                            prefix_idx = i
                            break
                        char_count += 1

            print(f"[调试] 前缀位置: {prefix_idx}, 前缀内容: '{full_text[prefix_idx:prefix_idx+len(prefix_pattern)]}'")
            
            dynamic_value_start = prefix_idx + len(prefix_pattern)
            
            if suffix_pattern:
                suffix_variants = self._get_punctuation_variants(suffix_pattern)
                print(f"[调试] 后缀变体: {suffix_variants}")
                suffix_idx = -1
                matched_suffix = None

                for variant in suffix_variants:
                    idx = full_text.find(variant, dynamic_value_start)
                    if idx != -1 and (suffix_idx == -1 or idx < suffix_idx):
                        suffix_idx = idx
                        matched_suffix = variant

                if suffix_idx == -1:
                    # 尝试标准化匹配（移除所有空白字符）
                    text_after_prefix = full_text[dynamic_value_start:]
                    normalized_after = normalize_text(text_after_prefix)
                    for variant in suffix_variants:
                        normalized_variant = normalize_text(variant)
                        idx_in_normalized = normalized_after.find(normalized_variant)
                        if idx_in_normalized != -1:
                            # 将标准化索引映射回原始文本索引
                            char_count = 0
                            for i, char in enumerate(text_after_prefix):
                                if not char.isspace():
                                    if char_count == idx_in_normalized:
                                        potential_idx = dynamic_value_start + i
                                        if suffix_idx == -1 or potential_idx < suffix_idx:
                                            suffix_idx = potential_idx
                                            matched_suffix = variant
                                        break
                                    char_count += 1

                if suffix_idx == -1:
                    print(f"警告: 在段落{para_idx}中未找到后缀 '{suffix_pattern}'，将标记到段落末尾")
                    dynamic_value_end = len(full_text)
                else:
                    print(f"[调试] 后缀位置: {suffix_idx}, 匹配的后缀: '{matched_suffix}'")
                    dynamic_value_end = suffix_idx
            else:
                # 没有后缀时，尝试自动检测常见的分隔符
                common_delimiters = ['，', ',', '。', '.', '；', ';', '、', '\n']
                dynamic_value_end = len(full_text)
                for delimiter in common_delimiters:
                    idx = full_text.find(delimiter, dynamic_value_start)
                    if idx != -1:
                        dynamic_value_end = idx
                        print(f"[调试] 自动检测到分隔符 '{delimiter}' 在位置 {idx}")
                        break

            dynamic_value = full_text[dynamic_value_start:dynamic_value_end].strip()
            if not dynamic_value:
                print(f"警告: 动态值为空，前缀 '{prefix_pattern}', 原始文本: '{full_text[dynamic_value_start:dynamic_value_end]}'")
                continue

            print(f"识别到动态值: '{dynamic_value}' (长度: {len(dynamic_value)})")

            # 使用isolate_run精确隔离动态值文本，然后应用高亮
            try:
                target_element = isolate_run(paragraph, dynamic_value_start, dynamic_value_end)
                # 应用背景高亮
                rPr = target_element.get_or_add_rPr()
                for shd in rPr.findall(qn("w:shd")):
                    rPr.remove(shd)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), color_hex)
                rPr.append(shd)
                print(f"已标记动态值: '{dynamic_value}'")
                return True
            except Exception as e:
                print(f"标记动态值时出错: {e}")
                continue
        
        print(f"警告: 在条款段落中未找到前缀 '{prefix_pattern}'")
        return False
    
    def _mark_clause_title(
        self,
        doc: Document,
        clause_paragraph_indices: List[int],
        clause_pattern: str,
        color_hex: str
    ) -> bool:
        """
        标记条款标题（兜底策略）
        
        当在条款中找不到动态值时，标红条款标题本身
        
        Args:
            doc: Document对象
            clause_paragraph_indices: 条款段落索引列表
            clause_pattern: 条款编号模式（如"第五条"）
            color_hex: 十六进制颜色值
        
        Returns:
            是否成功标记
        """
        clause_regex = re.compile(clause_pattern)
        
        for para_idx in clause_paragraph_indices:
            paragraph = doc.paragraphs[para_idx]
            para_text = paragraph.text
            
            if not clause_regex.search(para_text):
                continue
            
            print(f"[兜底] 在段落 {para_idx} 中找到条款标题 '{clause_pattern}'，将标红")
            
            full_text = ""
            run_positions = []
            
            for run in paragraph.runs:
                start_pos = len(full_text)
                full_text += run.text
                end_pos = len(full_text)
                run_positions.append((start_pos, end_pos, run))
            
            match = clause_regex.search(full_text)
            if not match:
                continue
            
            title_start = match.start()
            title_end = match.end()

            # 使用isolate_run精确隔离条款标题文本，然后应用高亮
            try:
                target_element = isolate_run(paragraph, title_start, title_end)
                # 应用背景高亮
                rPr = target_element.get_or_add_rPr()
                for shd in rPr.findall(qn("w:shd")):
                    rPr.remove(shd)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), color_hex)
                rPr.append(shd)
                print(f"[兜底] 已标记条款标题: '{clause_pattern}'")
                return True
            except Exception as e:
                print(f"标记条款标题时出错: {e}")
                continue
        
        print(f"[兜底] 警告: 未找到条款标题 '{clause_pattern}'")
        return False

    def mark_contract_with_rules(
        self,
        doc_path: str,
        marking_rules: List[Union[Dict[str, Any], 'MarkingRule']],
        output_path: str = None,
        add_comments: bool = True
    ) -> str:
        """
        根据标记规则自动识别并标记合同中的动态值（支持新格式）
        
        Args:
            doc_path: Word文档路径
            marking_rules: 标记规则列表，可以是MarkingRule对象或字典
            output_path: 输出文档路径，如果为None则覆盖原文件
            add_comments: 是否添加注释到文档末尾
        
        Returns:
            输出文档的路径
        """
        # 参数验证：确保文档路径存在
        if not os.path.exists(doc_path):
            raise IOError(f"文档路径不存在: {doc_path}")
        
        # 参数验证：确保标记规则不为空
        if not marking_rules:
            raise ValueError("标记规则不能为空")
        
        try:
            # 加载文档
            doc = Document(doc_path)
            
            # 收集所有注释信息
            comments_data = []
            
            # 按条款分组规则，避免重复处理同一条款
            clause_rules = defaultdict(list)
            
            for rule in marking_rules:
                # 统一转换为MarkingRule对象
                if isinstance(rule, dict):
                    rule = MarkingRule(**rule)
                
                clause_rules[rule.clause_pattern].append(rule)
            
            # 逐个条款处理
            for clause_pattern, rules in clause_rules.items():
                # 查找条款所在的段落索引
                clause_paragraph_indices = self._find_clause_paragraphs(doc, clause_pattern)
                
                # 如果未找到条款，记录警告并跳过
                if not clause_paragraph_indices:
                    print(f"警告: 未找到条款 '{clause_pattern}'，跳过该条款的所有规则")
                    continue
                
                # 收集该条款的所有标记任务
                mark_tasks = []
                for rule in rules:
                    color_hex = f"{rule.color[0]:02x}{rule.color[1]:02x}{rule.color[2]:02x}"
                    
                    if rule.mark_clause_title:
                        # 标记条款标题
                        mark_tasks.append({
                            "type": "title",
                            "pattern": clause_pattern,
                            "color_hex": color_hex
                        })
                    else:
                        # 标记动态值
                        mark_tasks.append({
                            "type": "value",
                            "prefix": rule.prefix_pattern,
                            "suffix": rule.suffix_pattern,
                            "color_hex": color_hex
                        })
                    
                    # 收集注释信息
                    if rule.comment:
                        comments_data.append({
                            "field_name": rule.field_name,
                            "clause": clause_pattern,
                            "comment": rule.comment
                        })
                
                # 一次性处理该条款的所有标记
                self._mark_multiple_in_clause(doc, clause_paragraph_indices, mark_tasks)
            
            # 如果需要添加注释到文档末尾
            if add_comments and comments_data:
                self._add_comments_to_document(doc, comments_data)
            
            # 确定保存路径
            save_path = output_path if output_path else doc_path
            
            # 保存文档
            doc.save(save_path)
            print(f"文档已保存至: {save_path}")
            return save_path
            
        except Exception as e:
            raise IOError(f"处理文档时出错: {str(e)}")
    
    def _mark_multiple_in_clause(
        self,
        doc: Document,
        clause_paragraph_indices: List[int],
        mark_tasks: List[Dict[str, Any]]
    ) -> bool:
        """
        在指定条款段落中一次性标记多个内容（使用isolate_run精确标记，保留原始格式）
        
        Args:
            doc: Document对象
            clause_paragraph_indices: 条款段落索引列表
            mark_tasks: 标记任务列表，每个任务包含type、pattern/color_hex等信息
        
        Returns:
            是否成功标记
        """
        import re
        
        # 标准化文本用于匹配
        def normalize_text(text):
            return ''.join(text.split())
        
        # 应用高亮样式到run
        def apply_highlight_to_run(run, color_hex):
            """应用背景高亮样式到run"""
            try:
                rPr = run._r.get_or_add_rPr()
                
                # 删除现有的阴影元素
                for shd in rPr.findall(qn("w:shd")):
                    rPr.remove(shd)
                
                # 创建新的阴影元素
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), color_hex)
                rPr.append(shd)
                return True
            except Exception as e:
                print(f"应用高亮样式时出错: {str(e)}")
                return False
        
        # 收集所有需要标记的区间
        all_mark_ranges = []
        
        for para_idx in clause_paragraph_indices:
            paragraph = doc.paragraphs[para_idx]
            para_text = paragraph.text
            
            full_text = ""
            for run in paragraph.runs:
                full_text += run.text
            
            for task in mark_tasks:
                task_type = task.get("type")
                color_hex = task.get("color_hex")
                
                if task_type == "title":
                    # 标记条款标题
                    pattern = task.get("pattern")
                    clause_regex = re.compile(pattern)
                    match = clause_regex.search(full_text)
                    if match:
                        all_mark_ranges.append({
                            "para_idx": para_idx,
                            "start": match.start(),
                            "end": match.end(),
                            "color_hex": color_hex
                        })
                        print(f"[标记标题] 段落 {para_idx}: '{pattern}'")
                
                elif task_type == "value":
                    # 标记动态值
                    prefix = task.get("prefix")
                    suffix = task.get("suffix")
                    
                    if not prefix:
                        continue
                    
                    # 查找前缀位置
                    prefix_idx = full_text.find(prefix)
                    if prefix_idx == -1:
                        normalized_full = normalize_text(full_text)
                        normalized_prefix = normalize_text(prefix)
                        normalized_idx = normalized_full.find(normalized_prefix)
                        if normalized_idx == -1:
                            continue
                        # 将标准化索引映射回原始文本索引
                        prefix_idx = 0
                        char_count = 0
                        for i, char in enumerate(full_text):
                            if not char.isspace():
                                if char_count == normalized_idx:
                                    prefix_idx = i
                                    break
                                char_count += 1
                    
                    dynamic_value_start = prefix_idx + len(prefix)
                    
                    # 确定动态值结束位置
                    if suffix:
                        suffix_idx = full_text.find(suffix, dynamic_value_start)
                        if suffix_idx == -1:
                            # 尝试标准化匹配
                            text_after_prefix = full_text[dynamic_value_start:]
                            normalized_after = normalize_text(text_after_prefix)
                            normalized_suffix = normalize_text(suffix)
                            idx_in_normalized = normalized_after.find(normalized_suffix)
                            if idx_in_normalized != -1:
                                char_count = 0
                                for i, char in enumerate(text_after_prefix):
                                    if not char.isspace():
                                        if char_count == idx_in_normalized:
                                            suffix_idx = dynamic_value_start + i
                                            break
                                        char_count += 1
                        
                        if suffix_idx == -1:
                            dynamic_value_end = len(full_text)
                        else:
                            dynamic_value_end = suffix_idx
                    else:
                        # 没有后缀时，尝试自动检测常见的分隔符
                        common_delimiters = ['，', ',', '。', '.', '；', ';', '、', '\n', ' ']
                        dynamic_value_end = -1
                        for delimiter in common_delimiters:
                            idx = full_text.find(delimiter, dynamic_value_start)
                            if idx != -1:
                                dynamic_value_end = idx
                                break
                        
                        # 如果没找到分隔符，尝试找下一个前缀（下一个字段的开始）
                        if dynamic_value_end == -1:
                            # 查找所有其他任务的前缀位置
                            next_prefix_start = len(full_text)
                            for other_task in mark_tasks:
                                if other_task.get("type") == "value" and other_task.get("prefix"):
                                    other_prefix = other_task.get("prefix")
                                    if other_prefix != prefix:  # 不是当前前缀
                                        other_idx = full_text.find(other_prefix, dynamic_value_start)
                                        if other_idx != -1 and other_idx < next_prefix_start:
                                            next_prefix_start = other_idx
                            
                            if next_prefix_start < len(full_text):
                                dynamic_value_end = next_prefix_start
                            else:
                                # 真的找不到结束位置，跳过此字段
                                print(f"[警告] 段落 {para_idx}: 无法确定字段 '{prefix}' 的结束位置，跳过")
                                continue
                    
                    dynamic_value = full_text[dynamic_value_start:dynamic_value_end].strip()
                    if dynamic_value and len(dynamic_value) < 200:  # 限制长度，防止标记过多内容
                        all_mark_ranges.append({
                            "para_idx": para_idx,
                            "start": dynamic_value_start,
                            "end": dynamic_value_end,
                            "color_hex": color_hex
                        })
                        print(f"[标记字段] 段落 {para_idx}: '{dynamic_value}'")
        
        if not all_mark_ranges:
            return False
        
        # 按段落分组标记区间
        para_mark_ranges = defaultdict(list)
        for mark_range in all_mark_ranges:
            para_mark_ranges[mark_range["para_idx"]].append(mark_range)
        
        # 逐个段落应用标记（使用isolate_run精确标记）
        for para_idx, ranges in para_mark_ranges.items():
            paragraph = doc.paragraphs[para_idx]
            
            # 合并所有标记区间（去重并排序）
            # 先按起始位置排序
            ranges.sort(key=lambda x: (x["start"], x["end"]))
            merged_ranges = []
            for r in ranges:
                if not merged_ranges:
                    merged_ranges.append(r.copy())
                elif r["start"] <= merged_ranges[-1]["end"]:
                    # 有重叠或相邻，合并区间
                    merged_ranges[-1]["end"] = max(merged_ranges[-1]["end"], r["end"])
                    # 如果新区间有颜色，优先使用新颜色（后定义的优先级更高）
                    if r.get("color_hex"):
                        merged_ranges[-1]["color_hex"] = r["color_hex"]
                else:
                    # 无重叠，添加新区间
                    merged_ranges.append(r.copy())
            
            # ⚠️ 关键：从后往前处理，这样前面的位置不会因为后面的修改而偏移
            for mark_range in reversed(merged_ranges):
                mark_start = mark_range["start"]
                mark_end = mark_range["end"]
                color_hex = mark_range["color_hex"]

                try:
                    # 使用isolate_run将目标文本隔离到独立run（保留原始格式）
                    target_run = isolate_run(paragraph, mark_start, mark_end)
                    # 应用高亮样式
                    apply_highlight_to_run(target_run, color_hex)
                except Exception as e:
                    print(f"标记失败 (段落 {para_idx}, 位置 {mark_start}-{mark_end}): {e}")
                    continue
        
        return True
    
    def _add_run_with_font(self, paragraph, text, run_positions, color_hex=None):
        """
        添加run并尝试保留原始字体格式
        
        Args:
            paragraph: 段落对象
            text: 文本内容
            run_positions: 原始run位置信息
            color_hex: 可选的颜色值
        """
        new_run = paragraph.add_run(text)
        
        # 尝试从原始run中继承字体格式
        for start_pos, end_pos, orig_run in run_positions:
            # 如果文本与原始run有重叠
            if orig_run.text and text in orig_run.text:
                if orig_run.font:
                    if orig_run.font.name:
                        new_run.font.name = orig_run.font.name
                    if orig_run.font.size:
                        new_run.font.size = orig_run.font.size
                    if orig_run.font.bold is not None:
                        new_run.font.bold = orig_run.font.bold
                    if orig_run.font.italic is not None:
                        new_run.font.italic = orig_run.font.italic
                    if orig_run.font.underline is not None:
                        new_run.font.underline = orig_run.font.underline
                break
        
        # 应用颜色
        if color_hex:
            new_run.font.color.rgb = RGBColor.from_string(color_hex)
    
    def _add_comments_to_document(self, doc: Document, comments_data: List[Dict[str, str]]):
        """
        在文档末尾添加注释说明
        
        Args:
            doc: Document对象
            comments_data: 注释数据列表
        """
        # 添加空行
        doc.add_paragraph()
        
        # 添加分隔线
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("=" * 50)
        separator_run.font.size = Pt(10)
        separator_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("【字段标注说明】")
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 添加每个字段的注释
        for item in comments_data:
            field_name = item.get("field_name", "")
            clause = item.get("clause", "")
            comment = item.get("comment", "")
            
            comment_para = doc.add_paragraph()
            
            # 字段名（红色加粗）
            field_run = comment_para.add_run(f"• {field_name}")
            field_run.font.size = Pt(10)
            field_run.font.bold = True
            field_run.font.color.rgb = RGBColor(255, 0, 0)
            
            # 条款信息
            clause_run = comment_para.add_run(f"（{clause}）")
            clause_run.font.size = Pt(9)
            clause_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # 注释内容
            if comment:
                desc_run = comment_para.add_run(f"：{comment}")
                desc_run.font.size = Pt(10)
                desc_run.font.color.rgb = RGBColor(0, 0, 0)

class MarkingRule:
    """标记规则配置类 - 更灵活的配置方式"""
    
    def __init__(
        self,
        clause_pattern: str,
        field_name: str,
        prefix_pattern: str = None,
        suffix_pattern: str = None,
        color: Tuple[int, int, int] = (255, 0, 0),
        comment: str = None,
        mark_clause_title: bool = False
    ):
        """
        初始化标记规则
        
        Args:
            clause_pattern: 条款编号模式（如"第五条"）
            field_name: 字段名称（用于标识和注释）
            prefix_pattern: 动态值前缀，如果为None且mark_clause_title为True，则标记条款标题
            suffix_pattern: 动态值后缀（可选）
            color: RGB颜色元组，默认为红色
            comment: 要添加的注释内容（可选）
            mark_clause_title: 是否标记条款标题本身（如第五条中的"第五条"）
        """
        self.clause_pattern = clause_pattern
        self.field_name = field_name
        self.prefix_pattern = prefix_pattern
        self.suffix_pattern = suffix_pattern
        self.color = color
        self.comment = comment
        self.mark_clause_title = mark_clause_title
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式（兼容旧方法）"""
        return {
            "clause_pattern": self.clause_pattern,
            "prefix_pattern": self.prefix_pattern,
            "suffix_pattern": self.suffix_pattern,
            "color": self.color,
            "field_name": self.field_name,
            "comment": self.comment,
            "mark_clause_title": self.mark_clause_title
        }


class ContractFieldMarker:
    """合同字段标记器 - 专门用于标记合同中的关键字段"""
    
    # 预定义的常用字段模式
    COMMON_FIELDS = {
        "不动产单元代码": {
            "prefix_pattern": "不动产单元代码为",
            "suffix_pattern": ",",
            "comment": "唯一标识该不动产单元的代码"
        },
        "宗地坐落": {
            "prefix_pattern": "宗地坐落于",
            "suffix_pattern": None,
            "comment": "宗地的具体地理位置"
        },
        "平面界址": {
            "prefix_pattern": "平面界址为",
            "suffix_pattern": None,
            "comment": "宗地的平面边界范围"
        },
        "用途": {
            "prefix_pattern": "用途为",
            "suffix_pattern": "。",
            "comment": "土地的使用用途"
        },
        "宗地总面积": {
            "prefix_pattern": "宗地总面积为大写",
            "suffix_pattern": "平方米",
            "comment": "宗地的总面积"
        },
        "出让面积": {
            "prefix_pattern": "出让面积为大写",
            "suffix_pattern": "平方米",
            "comment": "本次出让的土地面积"
        },
        "出让价款": {
            "prefix_pattern": "出让价款为大写",
            "suffix_pattern": "元",
            "comment": "土地出让的总价款"
        },
        "定金": {
            "prefix_pattern": "定金为大写",
            "suffix_pattern": "元",
            "comment": "合同定金金额"
        }
    }
    
    @classmethod
    def create_field_rule(
        cls,
        clause_pattern: str,
        field_name: str,
        color: Tuple[int, int, int] = (255, 0, 0),
        comment: str = None,
        **kwargs
    ) -> MarkingRule:
        """
        创建字段标记规则
        
        Args:
            clause_pattern: 条款编号（如"第五条"）
            field_name: 字段名称，可以是预定义名称或自定义
            color: 标记颜色
            comment: 自定义注释，如果为None则使用预定义注释
            **kwargs: 其他自定义参数
        
        Returns:
            MarkingRule对象
        """
        # 如果是预定义字段，使用预定义的模式
        if field_name in cls.COMMON_FIELDS:
            field_config = cls.COMMON_FIELDS[field_name].copy()
            # 如果传入了自定义comment，覆盖预定义的
            if comment:
                field_config["comment"] = comment
            # 更新其他自定义参数
            field_config.update(kwargs)
            return MarkingRule(
                clause_pattern=clause_pattern,
                field_name=field_name,
                color=color,
                **field_config
            )
        
        # 自定义字段
        return MarkingRule(
            clause_pattern=clause_pattern,
            field_name=field_name,
            color=color,
            comment=comment,
            **kwargs
        )
    
    @classmethod
    def create_clause_title_rule(
        cls,
        clause_pattern: str,
        color: Tuple[int, int, int] = (255, 0, 0),
        comment: str = None
    ) -> MarkingRule:
        """
        创建条款标题标记规则（标记"第五条"中的"第五条"）
        
        Args:
            clause_pattern: 条款编号（如"第五条"）
            color: 标记颜色
            comment: 注释内容
        
        Returns:
            MarkingRule对象
        """
        return MarkingRule(
            clause_pattern=clause_pattern,
            field_name=f"{clause_pattern}标题",
            prefix_pattern=None,
            suffix_pattern=None,
            color=color,
            comment=comment or f"{clause_pattern}条款标题",
            mark_clause_title=True
        )


if __name__ == "__main__":
    word_processor = WordProcessor()

    # 新的配置方式示例 - 更简洁、更灵活
    
    # 方式1: 使用 ContractFieldMarker 创建规则（推荐）
    marking_rules_new = [
        ContractFieldMarker.create_field_rule(
            clause_pattern="第五条",
            field_name="宗地总面积",
            
            comment="本条款为宗地基本信息条款",
            color=(255, 0, 0)
        ),
        # 标记第五条中的"第五条"标题本身
        ContractFieldMarker.create_clause_title_rule(
            clause_pattern="第五条",
            color=(255, 0, 0),
            comment="本条款为宗地基本信息条款"
        ),
        
        # 标记第五条中的各个字段
        ContractFieldMarker.create_field_rule(
            clause_pattern="第五条",
            field_name="不动产单元代码",
            color=(255, 0, 0)
        ),
        ContractFieldMarker.create_field_rule(
            clause_pattern="第五条",
            field_name="宗地坐落",
            color=(255, 0, 0)
        ),
        ContractFieldMarker.create_field_rule(
            clause_pattern="第五条",
            field_name="平面界址",
            color=(255, 0, 0)
        ),

        
        # 标记第六条中的字段
        ContractFieldMarker.create_field_rule(
            clause_pattern="第六条",
            field_name="用途",
            color=(255, 0, 0)
        ),
        
        # 方式2: 自定义字段（非预定义字段）
        MarkingRule(
            clause_pattern="第七条",
            field_name="出让面积",
            prefix_pattern="出让面积为大写",
            suffix_pattern="平方米",
            color=(255, 0, 0),
            comment="本次实际出让的土地面积"
        ),
    ]
    
    # 转换为旧格式兼容
    marking_rules_compat = [rule.to_dict() for rule in marking_rules_new]
    
    print("=== 新的标记规则配置 ===")
    for rule in marking_rules_new:
        print(f"字段: {rule.field_name}, 条款: {rule.clause_pattern}, "
              f"标记标题: {rule.mark_clause_title}, 注释: {rule.comment}")
    
    # 使用新的方法进行标记
    result = word_processor.mark_contract_with_rules(
        r"C:\Users\54299\Desktop\新版国有建设用地使用权出让合同.docx",
        marking_rules_new,
        r"C:\Users\54299\Desktop\新版国有建设用地使用权出让合同_marked.docx",
        add_comments=True  # 在文档末尾添加注释说明
    )
    print(f"\n标记完成，输出文件: {result}")
    
    # 也可以继续使用旧方法（兼容旧格式）
    # marking_rules_compat = [rule.to_dict() for rule in marking_rules_new]
    # result = word_processor.mark_contract_dynamic_values(
    #     r"C:\Users\54299\Desktop\新版国有建设用地使用权出让合同.docx",
    #     marking_rules_compat,
    #     r"C:\Users\54299\Desktop\新版国有建设用地使用权出让合同_marked.docx"
    # )
    # print(f"\n标记完成，输出文件: {result}")
